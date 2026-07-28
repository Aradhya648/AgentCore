"""Tests for deterministic Markdown → DOCX conversion."""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

import pytest
from docx import Document

from agentcore.docs_export.md_to_docx import (
    collect_image_srcs,
    convert_markdown_to_docx,
    docx_path_for_markdown,
    resolve_workspace_image_path,
)
import agentcore.docs_export.md_to_docx as md_to_docx_mod
from agentcore.docs_export.workspace_export import ExportMarkdownError, export_markdown_path
from agentcore.tools.sandbox import SubprocessSandbox
from agentcore.workspace.server import ServerWorkspace


def test_python_docx_is_lazy_until_convert(monkeypatch: pytest.MonkeyPatch) -> None:
    """Importing the module must not require python-docx until first convert.

    Sidecar mis-bundles historically crashed chat at tool registration via eager
    ``from docx import …``. Capability still ships in both runtimes; lazy load is
    import hygiene so a missing stack fails at convert, not at pipeline import.
    """
    monkeypatch.setattr(md_to_docx_mod, "_DocumentFactory", None)
    # Re-bind other sentinels so _ensure_docx runs a real import.
    for name in (
        "_WD_ALIGN_PARAGRAPH",
        "_RT",
        "_qn",
        "_OxmlElement",
        "_Cm",
        "_Inches",
        "_Pt",
        "_RGBColor",
        "_MAX_IMAGE_WIDTH",
    ):
        monkeypatch.setattr(md_to_docx_mod, name, None)
    assert md_to_docx_mod._DocumentFactory is None
    out = convert_markdown_to_docx("# hi")
    assert md_to_docx_mod._DocumentFactory is not None
    assert out.docx_bytes[:2] == b"PK"



def _tiny_png() -> bytes:
    from PIL import Image

    buf = io.BytesIO()
    Image.new("RGB", (8, 8), color=(20, 120, 200)).save(buf, format="PNG")
    return buf.getvalue()


_FIXTURE_MD = """# 标题一

## 小节二

这是一段含 [链接](https://example.com) 与 **加粗** 的正文。

- 无序甲
- 无序乙

1. 有序一
2. 有序二

```python
print("hello")
```

| 列A | 列B |
| --- | --- |
| 1 | 2 |

![示意图](./assets/chart.png)

![缺失](./assets/missing.png)
"""


def test_docx_path_for_markdown():
    assert docx_path_for_markdown("报告.md") == "报告.docx"
    assert docx_path_for_markdown("docs/a.markdown") == "docs/a.docx"
    assert docx_path_for_markdown("noext") == "noext.docx"


def test_collect_and_resolve_images():
    srcs = collect_image_srcs(_FIXTURE_MD)
    assert "./assets/chart.png" in srcs
    assert "./assets/missing.png" in srcs
    assert resolve_workspace_image_path("docs/报告.md", "./assets/chart.png") == (
        "docs/assets/chart.png"
    )
    assert resolve_workspace_image_path("报告.md", "https://x/y.png") is None


def test_convert_markdown_structure_and_missing_image_warning():
    png = _tiny_png()
    result = convert_markdown_to_docx(
        _FIXTURE_MD,
        images={
            "./assets/chart.png": png,
            "./assets/missing.png": None,
        },
    )
    assert result.docx_bytes[:2] == b"PK"
    assert any("缺图" in w and "missing.png" in w for w in result.warnings)

    # OOXML package is a zip; document.xml must exist.
    with zipfile.ZipFile(io.BytesIO(result.docx_bytes)) as zf:
        names = zf.namelist()
        assert "word/document.xml" in names
        xml = zf.read("word/document.xml").decode("utf-8")
        assert "标题一" in xml
        assert "hello" in xml
        assert "列A" in xml

    doc = Document(io.BytesIO(result.docx_bytes))
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert any("标题一" in t for t in texts)
    assert any("无序甲" in t for t in texts)
    assert any("有序一" in t for t in texts)
    assert len(doc.tables) >= 1
    assert doc.tables[0].cell(0, 0).text.strip() == "列A"


@pytest.mark.asyncio
async def test_export_markdown_path_writes_sibling(tmp_path: Path):
    root = tmp_path / "ws"
    root.mkdir()
    (root / "assets").mkdir()
    (root / "报告.md").write_text("# Hi\n\n![x](./assets/chart.png)\n", encoding="utf-8")
    (root / "assets" / "chart.png").write_bytes(_tiny_png())

    backend = ServerWorkspace(root=root, sandbox=SubprocessSandbox())
    out = await export_markdown_path(backend, "报告.md")
    assert out.output_path == "报告.docx"
    assert (root / "报告.docx").is_file()
    assert out.size_bytes == (root / "报告.docx").stat().st_size
    assert out.warnings == []


@pytest.mark.asyncio
async def test_export_markdown_path_missing_source(tmp_path: Path):
    backend = ServerWorkspace(root=tmp_path / "ws", sandbox=SubprocessSandbox())
    (tmp_path / "ws").mkdir()
    with pytest.raises(ExportMarkdownError, match="不存在"):
        await export_markdown_path(backend, "nope.md")
